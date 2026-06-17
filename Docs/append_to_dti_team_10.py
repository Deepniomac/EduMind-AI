from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
WORKING_DOC = ROOT / "DTI_TEAM_10_working.docx"
FINAL_DOC = ROOT / "DTI_TEAM_10_updated.docx"
ASSETS = ROOT / "team10_assets"

SCREENSHOTS = {
    "mindmap": Path(r"C:\Users\deepe\OneDrive\Pictures\Screenshots\Screenshot 2026-04-16 200851.png"),
    "recommendations": Path(r"C:\Users\deepe\OneDrive\Pictures\Screenshots\Screenshot 2026-04-16 201252.png"),
    "dashboard": Path(r"C:\Users\deepe\OneDrive\Pictures\Screenshots\Screenshot 2026-04-16 214046.png"),
    "cycle": Path(r"C:\Users\deepe\OneDrive\Pictures\Screenshots\Screenshot 2026-04-16 214113.png"),
}

NAVY = (10, 23, 38)
PANEL = (18, 34, 52)
OUTLINE = (53, 92, 132)
TEXT = (238, 245, 255)
MUTED = (165, 190, 225)
CYAN = (77, 205, 255)
YELLOW = (247, 202, 45)
GREEN = (55, 210, 135)
PINK = (255, 103, 136)


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_panel(draw: ImageDraw.ImageDraw, box, title: str, subtitle: str | None = None, fill=PANEL):
    draw.rounded_rectangle(box, radius=26, fill=fill, outline=OUTLINE, width=3)
    x1, y1, x2, _ = box
    draw.text((x1 + 24, y1 + 20), title, font=font(30, bold=True), fill=TEXT)
    if subtitle:
        draw.text((x1 + 24, y1 + 64), subtitle, font=font(18), fill=MUTED)


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color=CYAN, width=6):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        arrow = [(x2, y2), (x2 - 18 * direction, y2 - 10), (x2 - 18 * direction, y2 + 10)]
    else:
        direction = 1 if y2 > y1 else -1
        arrow = [(x2, y2), (x2 - 10, y2 - 18 * direction), (x2 + 10, y2 - 18 * direction)]
    draw.polygon(arrow, fill=color)


def create_dark_assets() -> dict[str, Path]:
    ensure_dir(ASSETS)

    architecture = ASSETS / "architecture_dark.png"
    workflow = ASSETS / "workflow_dark.png"
    palette = ASSETS / "frontend_palette.png"

    img = Image.new("RGB", (1600, 900), NAVY)
    draw = ImageDraw.Draw(img)
    draw.text((60, 36), "Figure 1. EduMind System Architecture", font=font(42, bold=True), fill=TEXT)
    draw_panel(draw, (80, 230, 310, 430), "Learner", "Browser")
    draw_panel(draw, (390, 180, 820, 470), "React + TypeScript Frontend", "Dashboard | Planner | Tutor | Quiz")
    draw_panel(draw, (900, 180, 1270, 470), "FastAPI Backend", "Validation | Routing | AI orchestration")
    draw_panel(draw, (1340, 230, 1530, 430), "Groq API", None, fill=(35, 44, 71))
    draw_arrow(draw, (310, 330), (390, 330))
    draw_arrow(draw, (820, 330), (900, 330))
    draw_arrow(draw, (1270, 330), (1340, 330))
    draw.text((420, 560), "Frontend preserves the adaptive learning experience while the backend handles AI requests securely.", font=font(24), fill=MUTED)
    img.save(architecture)

    img = Image.new("RGB", (1600, 920), NAVY)
    draw = ImageDraw.Draw(img)
    draw.text((60, 36), "Figure 2. Adaptive Learning Cycle", font=font(42, bold=True), fill=TEXT)
    steps = [
        ("Learn", CYAN),
        ("Test", (123, 218, 196)),
        ("Analyze", YELLOW),
        ("Adjust", (255, 162, 82)),
        ("Re-learn", GREEN),
    ]
    x = 85
    positions = []
    for label, color in steps:
        box = (x, 310, x + 240, 470)
        draw.rounded_rectangle(box, radius=30, fill=PANEL, outline=color, width=4)
        bbox = draw.textbbox((0, 0), label, font=font(30, bold=True))
        draw.text((x + (240 - (bbox[2] - bbox[0])) / 2, 365), label, font=font(30, bold=True), fill=TEXT)
        positions.append(box)
        x += 295
    for left, right in zip(positions, positions[1:]):
        draw_arrow(draw, (left[2], 390), (right[0], 390), color=YELLOW)
    draw.line([(1265, 470), (1265, 660), (205, 660), (205, 470)], fill=MUTED, width=5)
    draw_arrow(draw, (205, 470), (205, 471), color=MUTED)
    lines = [
        "Learn: explain concepts clearly and simply.",
        "Test: check concept understanding with quick questions.",
        "Analyze: identify weak areas and mistakes.",
        "Adjust: produce targeted revision guidance.",
        "Re-learn: reinforce concepts using the next study action.",
    ]
    y = 735
    for line in lines:
        draw.text((90, y), f"- {line}", font=font(24), fill=TEXT)
        y += 32
    img.save(workflow)

    img = Image.new("RGB", (1500, 420), NAVY)
    draw = ImageDraw.Draw(img)
    draw.text((50, 30), "Figure 3. EduMind Frontend Visual Language", font=font(40, bold=True), fill=TEXT)
    colors = [
        ("Deep navy", "#0b1626", (11, 22, 38)),
        ("Panel blue", "#132336", (19, 35, 54)),
        ("Outline blue", "#2d4f74", (45, 79, 116)),
        ("Accent cyan", "#4dcfff", (77, 207, 255)),
        ("Accent yellow", "#f7ca2d", (247, 202, 45)),
    ]
    x = 70
    for name, hex_value, rgb in colors:
        draw.rounded_rectangle((x, 120, x + 220, 240), radius=24, fill=rgb)
        draw.text((x, 270), name, font=font(24, bold=True), fill=TEXT)
        draw.text((x, 312), hex_value, font=font(22), fill=MUTED)
        x += 275
    img.save(palette)

    return {"architecture": architecture, "workflow": workflow, "palette": palette}


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    tc_pr.append(shade)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "D9EAF7")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(12)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(value)
            run.font.size = Pt(11)


def add_justified(document: Document, text: str, style: str = "Normal") -> None:
    para = document.add_paragraph(style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.2
    para.paragraph_format.space_after = Pt(6)
    para.add_run(text)


def add_centered(document: Document, text: str, style: str = "Normal", bold: bool = False) -> None:
    para = document.add_paragraph(style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold


def add_chapter(document: Document, number: int, title: str, intro_label: str | None = None) -> None:
    document.add_page_break()
    add_centered(document, f"CHAPTER {number}", bold=False)
    document.add_paragraph(title, style="Heading 1")
    if intro_label:
        document.add_paragraph(intro_label, style="Heading 3")


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        para = document.add_paragraph(style="List Paragraph")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_picture(document: Document, path: Path, caption: str, width: float = 6.4) -> None:
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(width))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True


def append_contents_and_body() -> Path:
    assets = create_dark_assets()
    document = Document(WORKING_DOC)

    document.add_paragraph("LIST OF DIAGRAMS", style="Heading 2")
    diagram_entries = [
        "Figure 1. EduMind System Architecture",
        "Figure 2. Adaptive Learning Cycle",
        "Figure 3. EduMind Frontend Visual Language",
        "Figure 4. EduMind Dashboard Overview",
        "Figure 5. Analyze Phase Priority and Cycle Tracker",
        "Figure 6. Mindmap Generator Interface",
        "Figure 7. AI Recommendations Interface",
    ]
    for entry in diagram_entries:
        add_justified(document, entry)

    add_chapter(document, 1, "Introduction", "INTRODUCTION")
    add_justified(
        document,
        "The rapid growth of digital education platforms has made learning resources more accessible than ever, but access alone does not guarantee understanding. "
        "Students often move between notes, videos, search tools, and generic chat systems without receiving structured help on what to revise, what they misunderstood, or what to do next."
    )
    add_justified(
        document,
        "EduMind is proposed as an AI-powered personalized learning and academic assistance platform that organizes the study process into a guided loop. "
        "Rather than behaving as a one-time answer engine, the platform supports concept explanation, practice, mistake analysis, targeted revision, and re-learning."
    )
    add_justified(
        document,
        "The project is important because it addresses a real gap in student learning behavior. Many learners need concise explanations, repeated reinforcement, progress visibility, and timely recommendations. "
        "EduMind brings these capabilities into a single professional academic platform designed around learner support."
    )
    document.add_paragraph("Project Objectives:", style="Heading 5")
    add_bullets(
        document,
        [
            "To provide students with a personalized AI-supported study assistant.",
            "To organize learning into a structured cycle of Learn, Test, Analyze, Adjust, and Re-learn.",
            "To help learners identify weak concepts and receive targeted revision support.",
            "To create an accessible web-based academic platform with a modern, student-friendly interface.",
        ],
    )

    add_chapter(document, 2, "Literature Review", "LITERATURE REVIEW")
    add_justified(
        document,
        "Adaptive learning research emphasizes that personalization improves student engagement, retention, and support quality when compared with static content delivery. "
        "Recent AI-assisted educational systems focus on learner modeling, targeted remediation, and context-sensitive guidance."
    )
    add_justified(
        document,
        "Existing tools such as standard note systems, fixed quiz portals, and chatbot interfaces each solve a part of the learning problem, but they often remain fragmented. "
        "Many systems explain a concept or test a student, yet they do not connect poor performance to a concrete revision pathway."
    )
    add_justified(
        document,
        "EduMind differs by combining personalized explanation, adaptive checkpoints, progress-focused dashboards, concept clustering, and revision-oriented recommendations into one study workflow. "
        "This makes the platform more suitable for continuous learning support than isolated educational utilities."
    )

    add_chapter(document, 3, "Design Thinking Process", "DESIGN THINKING PROCESS")
    document.add_paragraph("Empathize", style="Heading 5")
    add_justified(
        document,
        "The project began by identifying common student pain points such as uncertainty after reading a topic once, poor revision discipline, fragmented tool usage, and difficulty identifying weak concepts."
    )
    document.add_paragraph("Define", style="Heading 5")
    add_justified(
        document,
        "The core problem was defined as the absence of a learner-centered academic platform that can explain, evaluate, diagnose, and recommend the next study action in a single flow."
    )
    document.add_paragraph("Ideate", style="Heading 5")
    add_justified(
        document,
        "Different approaches were considered, including a chatbot-only tutor, a planner-only product, and a quiz-only system. "
        "The final concept combined these into a single adaptive workspace."
    )
    document.add_paragraph("Prototype", style="Heading 5")
    add_justified(
        document,
        "The prototype was implemented as a full-stack web application with an AI tutor, dashboard, mindmap generator, flashcards, adaptive quiz, and planner-style recommendation views."
    )
    document.add_paragraph("Test", style="Heading 5")
    add_justified(
        document,
        "Testing focused on whether the workflow remained understandable and helpful across multiple study screens, and whether the visual design reinforced clarity, focus, and progress awareness."
    )
    add_picture(document, assets["workflow"], "Figure 2. Adaptive learning cycle used in EduMind.", width=6.2)

    add_chapter(document, 4, "Proposed Solution", "PROPOSED SOLUTION")
    add_justified(
        document,
        "EduMind is a browser-based academic platform that combines AI tutoring and structured study management. "
        "The solution is designed to serve students who need more than static notes or generic answers. It helps them understand topics, validate learning, and re-enter weak concepts through actionable recovery steps."
    )
    add_justified(
        document,
        "The current solution includes a role-based study dashboard, AI tutor, adaptive quiz flow, flashcard review support, planner recommendations, and concept-map visualization. "
        "These modules work together to make the system feel like an academic workspace rather than a simple assistant window."
    )
    add_picture(document, assets["palette"], "Figure 3. Frontend color styling derived from the EduMind interface.", width=6.4)

    add_chapter(document, 5, "System Design", "SYSTEM DESIGN")
    add_justified(
        document,
        "The EduMind system follows a client-server architecture. The frontend is developed to manage learner interaction, navigation, visual feedback, and local study continuity, while the backend handles API requests, validation, and communication with the external AI model."
    )
    add_picture(document, assets["architecture"], "Figure 1. EduMind system architecture.", width=6.4)
    add_table(
        document,
        ["Module", "Description", "Role in the System"],
        [
            ["Dashboard", "Displays progress cards, recovery metrics, and current study phase", "Provides a learner-centric overview"],
            ["AI Tutor", "Explains topics and initiates guided study flow", "Supports concept learning"],
            ["Adaptive Quiz", "Checks understanding through targeted questions", "Measures concept retention"],
            ["Recommendations", "Suggests revision actions and priorities", "Supports recovery planning"],
            ["Mindmap Generator", "Visualizes topic clusters and concept relations", "Improves conceptual understanding"],
            ["Backend API", "Processes requests and AI communication", "Coordinates intelligent responses"],
        ],
        [1.2, 2.6, 2.4],
    )

    add_chapter(document, 6, "Implementation", "IMPLEMENTATION")
    add_justified(
        document,
        "The frontend of EduMind is built using React, TypeScript, Vite, and React Router. These technologies enable reusable UI components, maintainable state handling, and fast development workflows for an interactive educational interface."
    )
    add_justified(
        document,
        "The backend is implemented using Python and FastAPI, with request forwarding to an external Groq-compatible AI endpoint. "
        "The backend validates incoming study prompts, performs connectivity handling, and returns structured responses to the frontend."
    )
    add_justified(
        document,
        "A major implementation focus was preserving an academic yet modern user experience. The dark navy frontend theme, rounded card system, accent outlines, and cyan-to-yellow progress bars create a consistent visual identity that improves readability while preserving a professional feel."
    )
    add_picture(document, SCREENSHOTS["dashboard"], "Figure 4. EduMind dashboard overview screen.", width=6.6)
    add_picture(document, SCREENSHOTS["cycle"], "Figure 5. Analyze phase priorities and adaptive cycle tracker.", width=6.6)

    add_chapter(document, 7, "Testing and Results", "TESTING AND RESULTS")
    add_table(
        document,
        ["Test Case", "Testing Method", "Expected Result", "Outcome"],
        [
            ["Login and session flow", "Manual UI verification", "Learner workspace should load correctly", "Successful"],
            ["Dashboard rendering", "Component validation", "Metrics and phase-specific cards should appear", "Successful"],
            ["Tutor workflow", "End-to-end scenario test", "Explanation, quiz, and revision steps should connect", "Successful"],
            ["Mindmap generation", "Feature test using Data Structures topic", "Concept nodes should render visually", "Successful"],
            ["Recommendation display", "UI state verification", "Recovery suggestions should update clearly", "Successful"],
            ["Backend connectivity", "API route test", "AI mode should respond or fail safely", "Successful"],
        ],
        [1.5, 1.65, 2.5, 0.9],
    )
    add_justified(
        document,
        "The observed result is that EduMind successfully demonstrates the intended academic assistance workflow. "
        "The platform is especially strong in representing learner recovery, showing priorities, and converting concept revision into visible actions."
    )
    add_picture(document, SCREENSHOTS["mindmap"], "Figure 6. Mindmap generator screen showing concept-cluster rendering.", width=6.6)
    add_picture(document, SCREENSHOTS["recommendations"], "Figure 7. AI recommendations card with frontend progress styling.", width=4.7)

    add_chapter(document, 8, "Impact and Benefits", "IMPACT AND BENEFITS")
    document.add_paragraph("Social Benefits", style="Heading 5")
    add_justified(
        document,
        "EduMind can reduce confusion and study stress by giving students immediate learning support and clearer next actions. "
        "This helps learners who struggle with self-organization or repeated concept failure."
    )
    document.add_paragraph("Educational Benefits", style="Heading 5")
    add_justified(
        document,
        "The platform encourages active learning rather than passive reading. It improves concept reinforcement by linking explanations to practice, analysis, and re-learning."
    )
    document.add_paragraph("Technical Benefits", style="Heading 5")
    add_justified(
        document,
        "The project demonstrates the practical integration of AI-assisted tutoring, full-stack development, adaptive learning workflows, and visually guided progress design."
    )

    add_chapter(document, 9, "Future Enhancements", "FUTURE ENHANCEMENTS")
    add_bullets(
        document,
        [
            "Integrating real user authentication and database-backed learner profiles.",
            "Expanding support for more subjects, question types, and revision strategies.",
            "Adding analytics dashboards for faculty insight and classroom-level monitoring.",
            "Supporting PDF and note uploads for personalized concept extraction.",
            "Introducing multilingual learning support and speech-based interaction.",
            "Deploying the system for multi-user access in a hosted cloud environment.",
        ],
    )

    add_chapter(document, 10, "Conclusion", "CONCLUSION")
    add_justified(
        document,
        "EduMind presents a professional, student-centered academic assistance platform developed using a design thinking approach. "
        "It addresses the problem of fragmented and non-adaptive study tools by combining explanation, testing, analysis, recommendation, and re-learning into one integrated environment."
    )
    add_justified(
        document,
        "The system is not merely a chatbot but a structured educational workspace. Its combination of AI capabilities, modern full-stack implementation, and carefully styled academic interface gives it strong potential for future educational deployment."
    )

    document.add_page_break()
    document.add_paragraph("REFERENCES", style="Heading 1")
    refs = [
        "[1] H. Hariyanto, F. X. D. Kristianingsih, and R. Maharani, \"Artificial intelligence in adaptive education: a systematic review of techniques for personalized learning,\" Discover Education, 2025.",
        "[2] L. Fadieieva, \"Adaptive learning: a cluster-based literature review (2011-2022),\" Educational Technology Quarterly, 2023.",
        "[3] C. Merino-Campos, \"The Impact of Artificial Intelligence on Personalized Learning in Higher Education: A Systematic Review,\" Trends in Higher Education, 2025.",
        "[4] FastAPI Documentation. Available: https://fastapi.tiangolo.com/",
        "[5] React Documentation. Available: https://react.dev/",
        "[6] Vite Documentation. Available: https://vite.dev/guide/",
    ]
    for ref in refs:
        para = document.add_paragraph(style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.add_run(ref)

    document.save(FINAL_DOC)
    return FINAL_DOC


if __name__ == "__main__":
    print(append_contents_and_body())
