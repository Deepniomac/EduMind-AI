from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "generated_assets"
OUTPUT = ROOT / "EduMind_DTI_Final_Report.docx"

TITLE = "EduMind: AI-Powered Adaptive Study Assistant"
ACCENT = RGBColor(18, 63, 109)
ACCENT_LIGHT = RGBColor(226, 236, 248)
TEXT = RGBColor(33, 37, 41)
MUTED = RGBColor(90, 98, 104)


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    tc_pr.append(shade)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_doc_language(document: Document, language: str = "en-US") -> None:
    styles = document.styles
    for style_name in ("Normal", "Body Text", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = styles[style_name]
        except KeyError:
            continue
        rpr = style.element.get_or_add_rPr()
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:val"), language)


def set_default_style(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    p_format = normal.paragraph_format
    p_format.space_after = Pt(8)
    p_format.line_spacing = 1.2

    for style_name, size, color in (
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 11.5, RGBColor(50, 50, 50)),
    ):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    set_doc_language(document)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("EduMind DTI Final Report")
    hr.font.name = "Times New Roman"
    hr.font.size = Pt(9)
    hr.font.italic = True
    hr.font.color.rgb = MUTED


def add_paragraph(
    document: Document,
    text: str,
    *,
    bold_prefix: str | None = None,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    para = document.add_paragraph()
    para.alignment = alignment
    if bold_prefix and text.startswith(bold_prefix):
        run = para.add_run(bold_prefix)
        run.bold = True
        para.add_run(text[len(bold_prefix):])
    else:
        para.add_run(text)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        para = document.add_paragraph(style="List Paragraph")
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_section_heading(document: Document, number: int, title: str) -> None:
    document.add_paragraph(f"{number}. {title}", style="Heading 1")


def add_subheading(document: Document, title: str) -> None:
    document.add_paragraph(title, style="Heading 2")


def add_table(document: Document, headers: list[str], rows: list[list[str]], column_widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.width = Inches(column_widths[idx])
        set_cell_shading(cell, "D8E6F3")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(column_widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != len(row) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def try_font(size: int, bold: bool = False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
    ]
    for name in candidates:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, box, text, fill, outline=(38, 65, 108), text_fill=(24, 26, 27), radius=18):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    font = try_font(28, bold=True)
    lines = text.split("\n")
    total_h = 0
    metrics = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        metrics.append((line, bbox[2] - bbox[0], bbox[3] - bbox[1]))
        total_h += bbox[3] - bbox[1] + 6
    total_h -= 6
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, w, h in metrics:
        x = x1 + ((x2 - x1) - w) / 2
        draw.text((x, y), line, font=font, fill=text_fill)
        y += h + 6


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, fill=(38, 65, 108), width=6):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        arrow = [(x2, y2), (x2 - 20 * direction, y2 - 10), (x2 - 20 * direction, y2 + 10)]
    else:
        direction = 1 if y2 > y1 else -1
        arrow = [(x2, y2), (x2 - 10, y2 - 20 * direction), (x2 + 10, y2 - 20 * direction)]
    draw.polygon(arrow, fill=fill)


def create_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    title_font = try_font(40, bold=True)
    draw.text((60, 40), "Figure 1. High-Level Architecture of EduMind", font=title_font, fill=(18, 63, 109))

    draw_box(draw, (80, 220, 320, 420), "Learner\n(Browser)", (235, 244, 252))
    draw_box(draw, (390, 180, 840, 460), "React + TypeScript Frontend\nDashboard | Planner | Tutor\nQuiz | Session Context", (232, 247, 240))
    draw_box(draw, (910, 180, 1310, 460), "FastAPI Backend\nAPI Routing | Validation\nAI Bridge | Error Handling", (255, 243, 229))
    draw_box(draw, (1370, 220, 1550, 420), "Groq /\nAI API", (246, 235, 247))

    draw_arrow(draw, (320, 320), (390, 320))
    draw_arrow(draw, (840, 320), (910, 320))
    draw_arrow(draw, (1310, 320), (1370, 320))

    note_font = try_font(24)
    draw.text((520, 500), "Session state is preserved locally and adaptive prompts are routed through the backend.", font=note_font, fill=(70, 70, 70))
    img.save(path)


def create_flow_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 940), "white")
    draw = ImageDraw.Draw(img)
    title_font = try_font(40, bold=True)
    draw.text((60, 40), "Figure 2. Adaptive Learning Workflow", font=title_font, fill=(18, 63, 109))

    labels = ["Learn", "Test", "Analyze", "Adjust", "Re-learn"]
    fills = [
        (231, 243, 255),
        (235, 247, 240),
        (255, 248, 225),
        (255, 238, 230),
        (240, 235, 255),
    ]
    x = 80
    boxes = []
    for label, fill in zip(labels, fills):
        box = (x, 340, x + 250, 520)
        draw_box(draw, box, label, fill)
        boxes.append(box)
        x += 300

    for left, right in zip(boxes, boxes[1:]):
        draw_arrow(draw, (left[2], 430), (right[0], 430))

    loop_points = [(1380, 520), (1380, 690), (210, 690), (210, 520)]
    draw.line(loop_points, fill=(85, 85, 85), width=5)
    draw_arrow(draw, (210, 520), (210, 521), fill=(85, 85, 85), width=5)
    note_font = try_font(26)
    descriptions = [
        "Learn: explain the topic in simple language.",
        "Test: present short checkpoint questions.",
        "Analyze: identify what the learner missed.",
        "Adjust: create a focused revision plan.",
        "Re-learn: reinforce the weak concepts with a second pass.",
    ]
    y = 760
    for line in descriptions:
        draw.text((80, y), f"- {line}", font=note_font, fill=(70, 70, 70))
        y += 34

    img.save(path)


def create_design_thinking_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 950), "white")
    draw = ImageDraw.Draw(img)
    title_font = try_font(40, bold=True)
    draw.text((60, 40), "Figure 3. Design Thinking Applied to EduMind", font=title_font, fill=(18, 63, 109))

    center = (800, 470)
    radius = 280
    labels = [
        ("Empathize", (800, 150), (231, 243, 255)),
        ("Define", (1120, 320), (235, 247, 240)),
        ("Ideate", (1020, 710), (255, 248, 225)),
        ("Prototype", (580, 710), (255, 238, 230)),
        ("Test", (480, 320), (240, 235, 255)),
    ]
    for label, (x, y), fill in labels:
        draw_box(draw, (x - 120, y - 70, x + 120, y + 70), label, fill, radius=28)

    points = [(800, 220), (1070, 350), (980, 650), (620, 650), (530, 350)]
    for start, end in zip(points, points[1:] + [points[0]]):
        draw_arrow(draw, start, end)

    draw.ellipse((center[0] - 120, center[1] - 120, center[0] + 120, center[1] + 120), fill=(18, 63, 109))
    inner_font = try_font(28, bold=True)
    draw.text((720, 438), "Human-Centered\nInnovation", font=inner_font, fill="white")
    img.save(path)


def create_block_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 950), "white")
    draw = ImageDraw.Draw(img)
    title_font = try_font(40, bold=True)
    draw.text((60, 40), "Figure 4. Functional Block Diagram", font=title_font, fill=(18, 63, 109))

    draw_box(draw, (90, 240, 360, 420), "Input Layer\nLogin + Study Query", (231, 243, 255))
    draw_box(draw, (440, 160, 760, 330), "Learning Engine\nExplanation Generator", (235, 247, 240))
    draw_box(draw, (440, 380, 760, 550), "Assessment Engine\nQuiz + Answer Capture", (255, 248, 225))
    draw_box(draw, (860, 160, 1180, 330), "Analysis Engine\nWeakness Detection", (255, 238, 230))
    draw_box(draw, (860, 380, 1180, 550), "Recommendation Engine\nRevision Plan", (240, 235, 255))
    draw_box(draw, (1260, 240, 1510, 420), "Learner Output\nGuidance + Feedback", (230, 245, 237))

    draw_arrow(draw, (360, 330), (440, 245))
    draw_arrow(draw, (360, 330), (440, 465))
    draw_arrow(draw, (760, 245), (860, 245))
    draw_arrow(draw, (760, 465), (860, 465))
    draw_arrow(draw, (1180, 245), (1260, 330))
    draw_arrow(draw, (1180, 465), (1260, 330))
    img.save(path)


def build_figures() -> dict[str, Path]:
    ensure_dir(ASSET_DIR)
    paths = {
        "architecture": ASSET_DIR / "architecture.png",
        "workflow": ASSET_DIR / "workflow.png",
        "design_thinking": ASSET_DIR / "design_thinking.png",
        "block": ASSET_DIR / "block.png",
    }
    create_architecture_diagram(paths["architecture"])
    create_flow_diagram(paths["workflow"])
    create_design_thinking_diagram(paths["design_thinking"])
    create_block_diagram(paths["block"])
    return paths


def add_figure(document: Document, image_path: Path, caption: str, width: float = 6.6) -> None:
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Inches(width))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    cap.paragraph_format.space_after = Pt(10)


def title_page(document: Document) -> None:
    for _ in range(3):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.font.color.rgb = ACCENT

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Final DTI Project Report")
    sr.font.name = "Times New Roman"
    sr.font.size = Pt(15)
    sr.bold = True

    document.add_paragraph()
    info = document.add_table(rows=5, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    remove_table_borders(info)

    rows = [
        ("Project Title", TITLE),
        ("Team Members", "[Add team member names and roll numbers]"),
        ("Guide / Faculty", "[Add guide or faculty name]"),
        ("Department / College", "[Add department and college name]"),
        ("Academic Year", "2025-2026"),
    ]

    for row_idx, (label, value) in enumerate(rows):
        left, right = info.rows[row_idx].cells
        left.width = Inches(2.2)
        right.width = Inches(4.6)
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.name = "Times New Roman"
        rp = right.paragraphs[0]
        rr = rp.add_run(value)
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(12)
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for _ in range(5):
        document.add_paragraph()

    dept = document.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dept.add_run("Prepared for the Design Thinking and Innovation course")
    dr.font.name = "Times New Roman"
    dr.font.size = Pt(12)
    dr.italic = True

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("Replace the bracketed placeholders with your final submission details.")
    nr.font.name = "Times New Roman"
    nr.font.size = Pt(10.5)
    nr.font.color.rgb = MUTED

    document.add_page_break()


def create_report() -> Path:
    figures = build_figures()
    document = Document()
    set_default_style(document)
    configure_header_footer(document)
    title_page(document)

    add_section_heading(document, 2, "Abstract")
    add_paragraph(
        document,
        "EduMind is an AI-powered adaptive study assistant developed to improve the way students learn, test, and revise academic concepts. "
        "Traditional e-learning systems often provide static notes or one-time answers without checking whether a learner has actually understood the topic. "
        "EduMind addresses this limitation by organizing learning into a guided cycle of explanation, assessment, analysis, adjustment, and re-learning. "
        "The system uses a React and TypeScript frontend to provide an accessible study interface, while a FastAPI backend manages request handling and AI communication. "
        "The prototype currently demonstrates topic explanation, quiz generation, answer analysis, revision planning, dashboard-based progress visibility, and supporting pages such as planner and flashcard views. "
        "The project was shaped using a design thinking approach, beginning with student pain-point identification and moving through ideation, prototyping, and feedback-driven refinement. "
        "The resulting solution is a prototype that emphasizes personalized reinforcement rather than generic content delivery. "
        "EduMind shows how AI can be applied in a practical, student-centered learning workflow and offers a strong base for future expansion into richer analytics, broader subject support, and real-time adaptive study recommendations."
    )

    add_section_heading(document, 3, "Introduction")
    add_paragraph(
        document,
        "Education is rapidly moving toward digital and AI-supported environments, yet many learners still struggle with a familiar problem: they can access information easily, but they do not always receive the right guidance at the right time. "
        "Most conventional study tools focus on content delivery instead of concept reinforcement, error diagnosis, and structured revision."
    )
    add_paragraph(
        document,
        "EduMind was conceived as a response to this gap. The project focuses on making learning more adaptive by helping students move through a repeatable cycle: learn a concept, test understanding, analyze mistakes, adjust the revision plan, and re-learn the weak areas. "
        "This makes the platform suitable not only as a question-answer tool but also as a study companion that supports better academic habits."
    )
    add_paragraph(
        document,
        "The importance of this project lies in its human-centered focus. Many students need short explanations, immediate feedback, targeted revision prompts, and visible progress tracking. "
        "By combining these elements into one interface, EduMind aims to improve clarity, retention, and study discipline."
    )

    add_section_heading(document, 4, "Problem Statement")
    add_paragraph(
        document,
        "Students often rely on fragmented learning resources such as notes, videos, search engines, or generic chatbots. "
        "These tools may explain a topic, but they rarely verify understanding, pinpoint misconceptions, or guide the learner into a structured revision loop. "
        "As a result, learners may believe they understand a concept when they have only read it once. "
        "The core problem addressed by this project is the lack of an integrated, adaptive study system that can explain concepts, assess comprehension, analyze mistakes, and recommend the next best revision step in a simple and engaging workflow."
    )

    add_section_heading(document, 5, "Literature Review")
    add_paragraph(
        document,
        "Recent literature shows growing interest in adaptive and personalized learning systems. Hariyanto et al. [1] describe how AI-based adaptive education can personalize content delivery and learner support by using data-informed models and instructional adjustment. "
        "Fadieieva [2] reviews adaptive learning research between 2011 and 2022 and highlights the importance of tailoring content to learner needs, preferences, and performance patterns."
    )
    add_paragraph(
        document,
        "Merino-Campos [3] emphasizes that AI-supported personalized learning can improve learner engagement and flexibility in higher education, especially when systems adapt explanations and pacing to individual users. "
        "At the same time, design thinking literature from the Stanford d.school [4] shows that successful educational innovation should begin with empathy and iterative prototyping rather than technology-first decisions."
    )
    add_paragraph(
        document,
        "Existing study platforms provide valuable features such as digital flashcards, quiz practice, or note storage, but many of them are modular rather than truly integrated. "
        "EduMind differs by tying explanation, quiz-based checking, analysis, and revision planning into one loop. "
        "The project therefore builds on the strengths of prior adaptive learning research while keeping the implementation lightweight and directly aligned with student study behavior."
    )

    add_table(
        document,
        ["System / Approach", "Strengths", "Observed Limitations", "How EduMind Responds"],
        [
            [
                "Static e-learning content",
                "Easy access to notes and videos",
                "No feedback on whether the student understood the topic",
                "Adds testing and revision guidance after explanation",
            ],
            [
                "Generic chatbots",
                "Flexible question answering",
                "No built-in memory of study phase or recovery loop",
                "Introduces a structured Learn-Test-Analyze-Adjust-Re-learn cycle",
            ],
            [
                "Quiz-only tools",
                "Good for practice",
                "Often weak in concept explanation and remediation",
                "Connects quiz errors to targeted re-learning support",
            ],
        ],
        [1.65, 1.55, 2.25, 1.95],
    )

    add_section_heading(document, 6, "Design Thinking Process")
    add_paragraph(
        document,
        "The project followed the five-stage design thinking model of empathize, define, ideate, prototype, and test. "
        "This approach helped ensure that the solution stayed focused on learner needs instead of becoming only a technical demonstration."
    )
    add_figure(document, figures["design_thinking"], "Figure 3. Design thinking stages used to shape the EduMind prototype.")

    for title, text in [
        (
            "1. Empathize - Understanding Users",
            "The empathize stage focused on common difficulties faced by students during independent study: confusion after reading once, poor retention, lack of direction after making mistakes, and difficulty planning revision. "
            "Observations from everyday student behavior and the design goal of reducing repeated confusion informed the user profile."
        ),
        (
            "2. Define - Problem Definition",
            "The collected observations were synthesized into a clear challenge statement: students need a study assistant that does more than explain concepts once; it should verify understanding and recommend what to revise next. "
            "This definition narrowed the project toward adaptive support rather than general conversation."
        ),
        (
            "3. Ideate - Idea Generation",
            "During ideation, multiple options were considered, including a note summarizer, chatbot-only tutor, flashcard generator, and quiz portal. "
            "The most promising idea was a guided cycle that could combine tutoring, assessment, diagnosis, and revision support within one learning flow."
        ),
        (
            "4. Prototype - Model Creation",
            "The prototype was implemented as a web application with a React frontend and FastAPI backend. "
            "The interface includes a tutor cycle, progress-oriented dashboard, quiz page, study planner, and flashcard views to demonstrate how the learner experience can extend beyond a single prompt."
        ),
        (
            "5. Test - Evaluation of Solution",
            "Testing focused on validating whether the learner flow was understandable and complete: topic input, explanation, quiz answering, analysis generation, and revision guidance. "
            "Feedback from this stage informed refinements such as cleaner session handling, better workflow transitions, and clearer UI separation across pages."
        ),
    ]:
        add_subheading(document, title)
        add_paragraph(document, text)

    add_section_heading(document, 7, "Proposed Solution")
    add_paragraph(
        document,
        "EduMind is proposed as a browser-based adaptive study assistant that supports the learner from the first question to the revision stage. "
        "Instead of answering in isolation, the system is designed as a learning journey. A student begins by entering a study topic, receives a simple explanation, attempts a short quiz, and then gets feedback and a revision plan based on the mistakes made."
    )
    add_paragraph(
        document,
        "The current prototype demonstrates this idea through a tutor module and complementary screens such as the dashboard, planner, quiz page, flashcards, analytics-focused layout, and knowledge-view pages. "
        "Together these features communicate a broader vision: a learning system that not only teaches but also helps the learner recover weak concepts in a disciplined way."
    )
    add_bullets(
        document,
        [
            "Learner sign-in flow using demo user profiles.",
            "Topic explanation in either guided demo mode or backend AI mode.",
            "Quiz-based concept checking for immediate validation.",
            "Analysis of wrong answers and generation of a revision plan.",
            "Progress-oriented views such as dashboard, quiz map, planner, and flashcard reinforcement.",
        ],
    )
    add_paragraph(
        document,
        "For the final submitted report, this section can be strengthened further by inserting screenshots of the login page, dashboard, tutor flow, and quiz interface once you share the latest frontend captures."
    )

    add_section_heading(document, 8, "System Design")
    add_paragraph(
        document,
        "EduMind uses a client-server architecture in which the frontend is responsible for interaction, navigation, and local session continuity, while the backend handles API exposure, validation, and external AI communication. "
        "This separation keeps the system modular and easier to extend."
    )
    add_figure(document, figures["architecture"], "Figure 1. High-level architecture showing frontend, backend, and AI integration.")
    add_figure(document, figures["workflow"], "Figure 2. Adaptive learning workflow implemented by the tutor cycle.")
    add_figure(document, figures["block"], "Figure 4. Functional block diagram of EduMind modules.")
    add_paragraph(
        document,
        "The architecture diagram shows the data flow from the learner to the frontend, backend, and AI service. "
        "The workflow figure explains the pedagogical cycle, while the block diagram presents the main functional modules used to transform a study query into personalized feedback."
    )

    add_section_heading(document, 9, "Implementation")
    add_paragraph(
        document,
        "The frontend of EduMind is implemented using React, TypeScript, and Vite. React supports component-based UI composition, TypeScript improves reliability during refactoring, and Vite provides a fast local development workflow [6], [7]. "
        "Routing is handled through `react-router-dom`, and session persistence is supported through browser storage concepts aligned with the Web Storage API [8]."
    )
    add_paragraph(
        document,
        "The backend is implemented with Python and FastAPI. It exposes endpoints for tutor interaction and connectivity checks, loads environment variables, and forwards suitable prompts to the external Groq-compatible API. "
        "CORS middleware is enabled to support communication between the locally hosted frontend and backend [5]."
    )
    add_table(
        document,
        ["Component", "Technology Used", "Purpose"],
        [
            ["Frontend", "React + TypeScript + Vite", "Builds the study interface and page-level user experience"],
            ["Routing", "React Router", "Handles single-page navigation between learning modules"],
            ["Session Layer", "Context + localStorage", "Keeps learner state available between interactions"],
            ["Backend", "FastAPI + Python", "Processes requests and coordinates AI interaction"],
            ["External AI", "Groq API", "Generates dynamic study explanations in AI mode"],
        ],
        [1.35, 2.25, 2.9],
    )
    add_paragraph(
        document,
        "Deployment for the current prototype is local. The frontend runs through the Vite development server and the backend runs through Uvicorn. "
        "This environment is sufficient for functional testing and iterative improvement before a future hosted deployment."
    )

    add_section_heading(document, 10, "Testing and Results")
    add_paragraph(
        document,
        "Testing in this prototype focused on functional validation of the end-to-end learner workflow rather than large-scale classroom experimentation. "
        "The objective was to confirm that each core path works as expected: sign-in, topic entry, explanation rendering, quiz answering, analysis generation, revision planning, and backend connectivity checks."
    )
    add_table(
        document,
        ["Test Case", "Method", "Expected Outcome", "Observed Result"],
        [
            ["Demo sign-in", "Manual UI validation", "Learner profile loads successfully", "Passed"],
            ["Tutor start flow", "Scenario-based testing", "Question triggers explanation and quiz setup", "Passed"],
            ["Quiz completion", "Manual answer selection", "All questions can be answered and submitted", "Passed"],
            ["Analysis stage", "State transition check", "Mistakes produce review feedback and revision plan", "Passed"],
            ["Backend health check", "API endpoint check", "System detects connected or disconnected backend", "Passed"],
            ["API failure handling", "Negative-path validation", "Errors return clean HTTP failure responses", "Passed"],
        ],
        [1.55, 1.45, 2.8, 0.9],
    )
    add_paragraph(
        document,
        "The results indicate that the prototype successfully demonstrates the intended adaptive loop. "
        "The system is especially effective as a proof of concept for personalized revision guidance. "
        "However, the current results should be interpreted as prototype-level validation because the application still relies on demo content in several places and has not yet undergone broad user testing with real cohorts."
    )

    add_section_heading(document, 11, "Impact and Benefits")
    add_subheading(document, "Social Benefits")
    add_paragraph(
        document,
        "EduMind can support students who need more structured, repeated reinforcement than traditional note-based learning offers. "
        "By reducing confusion after the first explanation and by pointing learners toward their weak areas, the system can improve confidence and encourage better self-study habits."
    )
    add_subheading(document, "Economic Benefits")
    add_paragraph(
        document,
        "A browser-based study assistant can reduce dependence on multiple paid tools by combining explanation, practice, and revision support in one platform. "
        "For institutions, such systems may later reduce the time spent on repetitive doubt clarification while improving learner preparedness."
    )
    add_subheading(document, "Technical Improvement")
    add_paragraph(
        document,
        "The project demonstrates a technically meaningful combination of frontend state management, backend API orchestration, and AI-assisted tutoring. "
        "Its modular structure also makes it easier to extend into analytics, adaptive content sequencing, and persistent learner history."
    )

    add_section_heading(document, 12, "Future Enhancements")
    add_bullets(
        document,
        [
            "Add real user authentication and learner profiles instead of demo accounts.",
            "Store topic history, quiz attempts, and revision patterns in a database.",
            "Generate adaptive quizzes dynamically based on past errors and subject difficulty.",
            "Include richer analytics such as mastery trends, topic heat maps, and time-based revision alerts.",
            "Support image-based notes, PDF ingestion, voice interaction, and multilingual tutoring.",
            "Deploy the solution to a hosted environment for multi-user access and broader testing.",
        ],
    )

    add_section_heading(document, 13, "Conclusion")
    add_paragraph(
        document,
        "EduMind was developed as a design thinking-driven response to the limitations of static study tools and one-time chatbot interactions. "
        "The project identifies a real student need: learning support should not end with an explanation, but continue through assessment, diagnosis, and revision."
    )
    add_paragraph(
        document,
        "The final prototype demonstrates an integrated workflow built with React, TypeScript, FastAPI, and AI service connectivity. "
        "Its most important contribution is the structured learning cycle of Learn, Test, Analyze, Adjust, and Re-learn. "
        "Although the current system remains a prototype, it provides a credible foundation for a future adaptive learning platform with stronger personalization, analytics, and institutional relevance."
    )

    add_section_heading(document, 14, "References")
    references = [
        "[1] H. Hariyanto, F. X. D. Kristianingsih, and R. Maharani, \"Artificial intelligence in adaptive education: a systematic review of techniques for personalized learning,\" Discover Education, vol. 4, art. 458, 2025.",
        "[2] L. Fadieieva, \"Adaptive learning: a cluster-based literature review (2011-2022),\" Educational Technology Quarterly, no. 3, pp. 319-366, 2023. doi: 10.55056/etq.613.",
        "[3] C. Merino-Campos, \"The Impact of Artificial Intelligence on Personalized Learning in Higher Education: A Systematic Review,\" Trends in Higher Education, vol. 4, no. 2, art. 17, 2025. doi: 10.3390/higheredu4020017.",
        "[4] Stanford d.school, \"Design Thinking Bootleg,\" [Online]. Available: https://dschool.stanford.edu/resources/the-bootcamp-bootleg.",
        "[5] FastAPI Documentation, \"CORS (Cross-Origin Resource Sharing),\" [Online]. Available: https://fastapi.tiangolo.com/tutorial/cors/.",
        "[6] React Documentation, \"React Reference Overview,\" [Online]. Available: https://react.dev/reference/react.",
        "[7] Vite Documentation, \"Getting Started,\" [Online]. Available: https://vite.dev/guide/.",
        "[8] MDN Web Docs, \"Web Storage API,\" [Online]. Available: https://developer.mozilla.org/en-US/docs/Web/API/Web_Storage_API.",
    ]
    for ref in references:
        add_paragraph(document, ref, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    ensure_dir(ASSET_DIR)
    result = create_report()
    print(result)
